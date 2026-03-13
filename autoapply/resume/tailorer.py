import logging
import os
import shutil
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional

from ..config import config
from ..llm.analyzer import JobAnalyzer

logger = logging.getLogger(__name__)


class ResumeTailorer:
    """Tailor a resume for a specific job using LLM suggestions."""

    def __init__(self, analyzer: JobAnalyzer = None):
        from ..llm.analyzer import analyzer as default_analyzer
        self.analyzer = analyzer or default_analyzer

    def tailor_for_job(
        self,
        resume_path: str,
        job_title: str,
        job_description: str,
        output_dir: str = None,
    ) -> Optional[str]:
        """Create a tailored resume copy for a specific job."""
        if not os.path.exists(resume_path):
            logger.error(f"Resume not found: {resume_path}")
            return None

        ext = os.path.splitext(resume_path)[1].lower()
        if ext not in (".docx", ".pdf"):
            logger.error(f"Unsupported resume format: {ext}")
            return None

        # Get tailoring suggestions from LLM
        from ..resume.parser import parse_resume
        parsed = parse_resume(resume_path)
        suggestions = self.analyzer.tailor_resume(
            job_title=job_title,
            job_description=job_description,
            resume_summary=parsed.get("full_text", "")[:2000],
        )

        if not suggestions:
            logger.warning("No tailoring suggestions from LLM, using original resume")
            return resume_path

        # Only DOCX can be programmatically edited
        if ext == ".docx":
            return self._tailor_docx(resume_path, suggestions, output_dir)
        else:
            # For PDF, just copy and log suggestions
            output_path = self._copy_resume(resume_path, output_dir, f"tailored_{job_title}")
            logger.info(f"PDF resume copied (manual tailoring needed). Suggestions: {suggestions}")
            return output_path

    def _tailor_docx(
        self, resume_path: str, suggestions: Dict, output_dir: str = None
    ) -> Optional[str]:
        """Apply tailoring suggestions to a DOCX file."""
        try:
            from docx import Document

            doc = Document(resume_path)
            suggested_summary = suggestions.get("suggested_summary", "")
            keywords = suggestions.get("keywords_to_add", [])
            skills_to_highlight = suggestions.get("skills_to_highlight", [])

            # Replace summary paragraph if suggested
            if suggested_summary and doc.paragraphs:
                # Try to find and update the summary paragraph
                for para in doc.paragraphs[:10]:
                    text = para.text.strip()
                    if len(text) > 50 and not text.isupper():
                        # Likely the summary paragraph
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = suggested_summary
                        break

            # Add keywords section at the bottom if not present
            skills_text = ", ".join(skills_to_highlight + keywords)
            if skills_text:
                # Append to the skills section if it exists, otherwise add at end
                skills_added = False
                for para in doc.paragraphs:
                    if "skills" in para.text.lower() and len(para.text) < 30:
                        # Found skills header - find next paragraph and append
                        skills_added = True
                        break

                if not skills_added:
                    doc.add_paragraph(f"Key Skills: {skills_text}")

            # Save tailored resume
            output_path = self._get_output_path(resume_path, output_dir, "tailored")
            doc.save(output_path)
            logger.info(f"Tailored resume saved: {output_path}")
            return output_path

        except ImportError:
            logger.error("python-docx not installed")
            return resume_path
        except Exception as e:
            logger.error(f"DOCX tailoring error: {e}")
            return resume_path

    def _copy_resume(self, source: str, output_dir: str = None, prefix: str = "copy") -> str:
        """Copy resume to output directory."""
        output_path = self._get_output_path(source, output_dir, prefix)
        shutil.copy2(source, output_path)
        return output_path

    def _get_output_path(self, source: str, output_dir: str = None, prefix: str = "") -> str:
        """Generate output file path for tailored resume."""
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        ext = os.path.splitext(source)[1]
        basename = os.path.splitext(os.path.basename(source))[0]
        filename = f"{basename}_{prefix}_{timestamp}{ext}"

        if output_dir:
            Path(output_dir).mkdir(parents=True, exist_ok=True)
            return os.path.join(output_dir, filename)
        else:
            return os.path.join(os.path.dirname(source), filename)


tailorer = ResumeTailorer()
